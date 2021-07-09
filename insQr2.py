from docx import Document
from docx.shared import Inches
from os import listdir
from os.path import isfile, join
import os
import shutil
import safeRandom
import gen2qr
import yaml


def movefile(srcp, destp, name):
    src = srcp+'/'+name
    dst = destp+'/'+name
    shutil.move(src, dst)


def getFiles(fpath):
    return [f for f in listdir(fpath) if isfile(join(fpath, f))]


def popSheet6576(fpath, flist, NumberItems, ftemplate, fsave, moveDir='./usedDuoQr'):
    availableFiles = len(flist)

    if NumberItems > availableFiles:
        raise ValueError('Number of available files:'+str(availableFiles) +
                         ' is not enough for this run: Need '+str(NumberItems))

    document = Document(ftemplate)
    for i in range(0, NumberItems):
        rowIdx = i // 4
        cellIdx = (i % 4)*2
        nameOnly = flist.pop(0)
        filename = fpath + '/'+nameOnly
        document.tables[0].rows[rowIdx].cells[cellIdx].paragraphs[0].add_run()
        document.tables[0].rows[rowIdx].cells[cellIdx].paragraphs[0].runs[0].add_picture(
            filename, width=Inches(1.44))
        movefile(fpath, moveDir, nameOnly)

    document.save(fsave)
    shutil.move(fsave, f"{os.getcwd()}/QrSheets")


# left here for documentation
def test():
    flist = getFiles('./duoQr')
    popSheet6576('./duoQr', flist, 32,
                 'Avery6576DurableIDLabels.docx', 'duo2.docx')
    flist2 = getFiles('./duoQr')
    popSheet6576('./duoQr', flist2, 32,
                 'Avery6576DurableIDLabels.docx', 'duo3.docx')
    flist3 = getFiles('./duoQr')
    popSheet6576('./duoQr', flist3, 32,
                 'Avery6576DurableIDLabels.docx', 'duo4.docx')


def main():

    with open("qrCfg.yaml", 'r') as stream:
        cfg = yaml.load(stream)

    print(cfg)
    safeR = safeRandom.PropertyId(cfg['db'])
    numberPages = len(cfg['outputfiles'])
    # create property tag images to be used later
    # images are save in directory {cfg['qrDir']}
    gen2qr.createStickersForPage(
        numberPages*cfg['stickersPerPage'], safeR, cfg['qrDir'], cfg['numDigits'])
    safeR.close()

    # take created property images and insert them into a word document for printing
    for outputfile in cfg['outputfiles']:
        flist = getFiles(cfg['qrDir'])
        popSheet6576(cfg['qrDir'], flist, cfg['stickersPerPage'],
                     cfg['stickertemplate'], outputfile, cfg['usedqrDir'])


if __name__ == "__main__":
    main()
